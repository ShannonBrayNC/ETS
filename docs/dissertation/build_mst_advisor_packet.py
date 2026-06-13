from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
MD_PATH = OUT_DIR / "MST_ADVISOR_REVIEW_PACKET.md"
DOCX_PATH = OUT_DIR / "MST_ADVISOR_REVIEW_PACKET.docx"


TITLE = "Return-to-PhD Advisor Review Packet"
SUBTITLE = "Evidence Transparency Systems (ETS)"
AUTHOR = "Shannon Bray"
DATE = "Prepared June 2026"


SECTIONS: list[tuple[str, list[str]]] = [
    (
        "Cover Note",
        [
            "Purpose: This packet is prepared for advisor review as I evaluate returning to complete my PhD at Missouri University of Science and Technology. The research program is Evidence Transparency Systems (ETS), a formal architecture for verifiable digital evidence, bounded trust coordination, replayable audit artifacts, and governance traceability under adversarial and incomplete observation conditions.",
            "My immediate request is not for final dissertation approval. It is for faculty guidance on whether ETS can be shaped into a defensible PhD dissertation, what committee expectations would need to be met, what prior academic standing or administrative steps must be resolved, and what evidence would be required to resume doctoral work responsibly.",
            "The attached rough dissertation draft is a synthesis of the current ETS research corpus. It is intentionally conservative: ETS does not claim universal truth, perfect completeness, full Byzantine consensus, legal sufficiency, or production certification. It claims bounded verifiability of recorded digital evidence under explicit formal, cryptographic, computational, and governance assumptions.",
        ],
    ),
    (
        "Advisor Review Letter Draft",
        [
            "Dear Professor,",
            "I hope you are doing well. I am writing to ask whether you would be willing to review my current research direction and advise me on the possibility of returning to complete my PhD at Missouri S&T.",
            "Since my previous doctoral work, I have continued developing a research program called Evidence Transparency Systems (ETS). ETS addresses a verification gap in modern digital systems: evidence may exist in logs, AI audit trails, compliance records, observability systems, and operational reports, but independent parties often cannot verify integrity, ordering, omission risk, replayability, or provenance without trusting the same system under review.",
            "The research now includes a dissertation-style rough draft, a Sprint 2 claim audit, a Sprint 3 bibliography and related-work matrix, a Sprint 4 formal-methods audit and proof-status table, a Sprint 5 implementation/reproducibility audit, a Sprint 6 publication pipeline, a Sprint 7 dissertation assembly plan, and a Sprint 8 defense-preparation package. It also includes TLA+ model directions, proof and theorem registries, implementation traceability, reproducibility materials, and an evaluation plan. The central thesis is that ETS can provide a formal architecture for computationally bounded evidentiary coordination by combining canonical evidence records, append-only transparency logs, cryptographic proof artifacts, verifier federation, replay, confidence semantics, and reproducible validation.",
            "I would appreciate your candid assessment of whether this work can be brought into alignment with Missouri S&T doctoral expectations, what parts would need to be strengthened, and whether there may be a viable path to re-enter the PhD process. I am also exploring funding and veteran education benefit pathways, including VR&E, remaining GI Bill eligibility if applicable, Missouri veteran tuition reductions, and potential assistantship support.",
            "If you are open to it, I would be grateful for a short meeting to discuss the attached packet and determine the right next step.",
            "Respectfully,",
            "Shannon Bray",
        ],
    ),
    (
        "Research Prospectus Summary",
        [
            "Working title: Evidence Transparency Systems: A Formal Architecture for Computationally Bounded Evidentiary Coordination Under Adversarial and Incomplete Observation Conditions.",
            "Research problem: Modern digital systems increasingly generate consequential records about events, decisions, model outputs, compliance actions, and operational behavior. These records often remain controlled by the same institutions or platforms whose behavior is being evaluated. ETS addresses the resulting verification gap by asking what independent reviewers can prove about recorded evidence without requiring privileged trust in the originating system.",
            "Thesis statement: Evidence Transparency Systems can provide a defensible protocol architecture for verifiable digital evidence by combining canonical evidence records, append-only transparency logs, proof-carrying audit artifacts, verifier federation, and reproducible replay. ETS improves independent verifiability and governance traceability without claiming semantic truth, perfect completeness, full Byzantine consensus, or universal correctness.",
            "Core research questions: How can heterogeneous events be canonicalized so independent verifiers compute the same hashes? How can append-only transparency systems expose tampering, reordering, and forked histories? How can verifier federation detect divergence, stale state, omission suspicion, or inconsistent roots? How can asynchronous transport and partial visibility be modeled without overstating liveness? How should confidence and trust be represented as bounded semantics rather than absolute truth labels? How can protocol requirements trace to code, formal models, tests, and reproducible artifacts?",
            "Methodology: The dissertation would use design science and systems research methods. It defines protocol requirements, implements a reference system, models core properties formally, evaluates behavior through deterministic experiments, and compares ETS against transparency logs, distributed systems, formal verification, auditability, AI governance, and reproducible systems research.",
        ],
    ),
    (
        "Proposed Dissertation Contributions",
        [
            "Theoretical contribution: A bounded evidence-transparency model that separates evidence from truth, observation from certainty, confidence from proof, omission suspicion from completeness, and disagreement from failure.",
            "Formal-methods contribution: TLA+ and related formal models for append-only logs, verifier federation, liveness, replay, temporal adversarial behavior, and asynchronous transport under explicit assumption boundaries.",
            "Systems contribution: A protocol architecture using canonical evidence objects, cryptographic hashes, append-only logs, Merkle-style proof semantics, signed roots, proof bundles, verifier APIs, replay, and federation checks.",
            "Implementation contribution: A reference research implementation with tests, CLI/API paths, reproducibility harnesses, benchmark scaffolding, and traceability from claims to artifacts.",
            "Evaluation contribution: Deterministic experiments for fork detection, omission suspicion, replay mismatch, stale-state recovery, asynchronous transport, verifier convergence, and reproducible artifact packages.",
            "Governance contribution: Application patterns for AI governance, compliance audit, research integrity, and institutional accountability that preserve human review and do not convert technical proof into automatic legitimacy.",
        ],
    ),
    (
        "Existing Work Completed",
        [
            "A rough PhD dissertation draft has been prepared in the ETS repository: docs/dissertation/ETS_PHD_DISSERTATION_ROUGH_DRAFT.md and docs/dissertation/ETS_PHD_DISSERTATION_ROUGH_DRAFT.docx.",
            "Sprint 2 claim-discipline artifacts have been prepared: docs/dissertation/CLAIM_AUDIT.md, docs/dissertation/RESEARCH_ARTIFACT_MAP.md, and docs/dissertation/SPRINT_2_READINESS_REPORT.md.",
            "Sprint 3 literature-hardening artifacts have been prepared: docs/dissertation/BIBLIOGRAPHY.md, docs/dissertation/RELATED_WORK_MATRIX.md, docs/dissertation/SPRINT_3_READINESS_REPORT.md, and an expanded docs/dissertation/LITERATURE_REVIEW.md.",
            "Sprint 4 formal-methods artifacts have been prepared: docs/dissertation/FORMAL_METHODS_AUDIT.md, docs/dissertation/PROOF_STATUS_TABLE.md, docs/dissertation/MODEL_CHECKING_COMMAND_LOG.md, and docs/dissertation/SPRINT_4_READINESS_REPORT.md.",
            "Sprint 5 implementation and reproducibility artifacts have been prepared: docs/dissertation/IMPLEMENTATION_REPRODUCIBILITY_AUDIT.md, docs/dissertation/EXPERIMENT_ARTIFACT_PLAN.md, docs/dissertation/GOLDEN_VECTOR_COVERAGE.md, and docs/dissertation/SPRINT_5_READINESS_REPORT.md.",
            "Sprint 6 publication pipeline artifacts have been prepared: docs/dissertation/PAPER_PIPELINE_ROADMAP.md, docs/dissertation/PAPER_ABSTRACTS_AND_OUTLINES.md, docs/dissertation/PAPER_CLAIM_EVIDENCE_MAP.md, docs/dissertation/VENUE_STRATEGY.md, and docs/dissertation/SPRINT_6_READINESS_REPORT.md.",
            "Sprint 7 dissertation assembly artifacts have been prepared: docs/dissertation/DISSERTATION_ASSEMBLY_PLAN.md, docs/dissertation/CHAPTER_INTEGRATION_CHECKLIST.md, docs/dissertation/FIGURE_TABLE_PLAN.md, docs/dissertation/COMMITTEE_DRAFT_READINESS.md, and docs/dissertation/SPRINT_7_READINESS_REPORT.md.",
            "Sprint 8 defense-preparation artifacts have been prepared: docs/dissertation/DEFENSE_QA.md, docs/dissertation/DEFENSE_DECK_PLAN.md, docs/dissertation/ARTIFACT_WALKTHROUGH_SCRIPT.md, docs/dissertation/FINAL_REVISION_CHECKLIST.md, and docs/dissertation/SPRINT_8_READINESS_REPORT.md.",
            "The existing dissertation corpus includes prospectus material, dissertation structure, literature review, formal foundations, formal architecture, evidence theory, evaluation and benchmarks, reproducibility notes, theorem registries, proof indices, temporal liveness notes, probabilistic Byzantine convergence notes, symbolic verification notes, implementation traceability, and contribution summaries.",
            "The research corpus includes ETS research-paper release candidates, an executable research plan, asynchronous transport research, verifier federation and convergence notes, TLA execution and validation guidance, reproducibility appendix, formal theorems, and a formal traceability matrix.",
            "The formal artifact set includes TLA+ models for append-only logs, verifier federation, asynchronous transport, temporal Byzantine federation, probabilistic trust, liveness federation, and universal temporal liveness variants; Lean proof-development materials for temporal liveness, fairness, and Byzantine temporal properties; and test suites for benchmarks, experiments, federation, probabilistic behavior, governance, and dissertation deliverables.",
            "A dissertation deliverables test currently passes in the local repository: python -m pytest tests/unit/test_dissertation_deliverables.py -q.",
        ],
    ),
    (
        "Known Gaps Before Dissertation Readiness",
        [
            "Advisor and committee alignment: The research needs review by a Missouri S&T advisor to determine departmental fit, committee expectations, admissibility or reinstatement path, and whether the current artifacts can count toward a dissertation trajectory.",
            "Literature review: The draft needs a normalized academic bibliography and deeper engagement with primary literature in transparency logs, distributed systems, formal methods, auditability, AI governance, evidence theory, and reproducible systems research.",
            "Formal proof maturity: The theorem registry should be converted into a final proof-completion table that clearly separates proved, modeled, tested, pending, and not-claimed properties.",
            "Empirical results: Benchmark and experiment outputs need final scenario manifests, commands, environment summaries, generated artifacts, result tables, interpretation notes, and replication instructions.",
            "Cross-implementation validation: Canonicalization and proof-bundle semantics would be stronger with golden test vectors and, if feasible, a second implementation or independent verification script.",
            "Institutional formatting: The dissertation draft must be brought into Missouri S&T thesis/dissertation formatting and submission requirements after committee direction.",
        ],
    ),
    (
        "Funding and Return Path Questions",
        [
            "Question 1: What is my current academic status at Missouri S&T, and what formal steps are required to return to doctoral study if I was previously enrolled?",
            "Question 2: If I am eligible for VA disability-related education support, can VR&E Chapter 31 approve doctoral study as part of a suitable employment or rehabilitation plan?",
            "Question 3: Do I have remaining Post-9/11 GI Bill eligibility, and how would it interact with VR&E, assistantship funding, or tuition reduction programs?",
            "Question 4: If I qualify as a combat veteran under the Missouri Returning Heroes Act, can graduate tuition and fees be reduced to no more than 30% of normal cost for doctorate-level study, subject to eligibility and time limits?",
            "Question 5: Could a research or teaching assistantship apply, and would it include tuition remission, stipend support, or fee reduction?",
            "Question 6: Which office should coordinate certification and sequencing: advisor/department, Graduate Education, Student Financial Services, Military & Veterans Service Center, or VA VR&E counselor?",
        ],
    ),
    (
        "Potential Elon University Collaboration",
        [
            "Elon University in North Carolina may be a useful academic collaboration partner after the Missouri S&T path is clarified. The strongest fit is not degree administration, but applied research collaboration, undergraduate research, cybersecurity/data science modules, capstone projects, and AI governance or digital-evidence education.",
            "Possible collaboration theme: Evidence Transparency Systems for undergraduate research in cybersecurity, AI accountability, reproducible experiments, and verifiable digital evidence.",
            "Possible Elon-facing deliverables: a two-page collaboration brief, a student research module, an ETS lab exercise, a capstone project outline, and a joint workshop proposal. These should be prepared only after advisor feedback confirms the dissertation framing.",
        ],
    ),
    (
        "Proposed Next Meeting Agenda",
        [
            "1. Confirm whether ETS is a viable PhD research direction at Missouri S&T.",
            "2. Identify the correct department, advisor, committee, and administrative return path.",
            "3. Review the thesis statement, contribution claims, and non-claim boundaries.",
            "4. Decide which artifacts need to be converted into publishable papers.",
            "5. Establish a short list of experiments and formal proofs required for dissertation readiness.",
            "6. Discuss funding routes and whether assistantship support is possible.",
            "7. Decide whether Elon University collaboration should be pursued now or after re-entry approval.",
        ],
    ),
    (
        "Attachments and Repository Artifacts",
        [
            "Primary attachment: docs/dissertation/ETS_PHD_DISSERTATION_ROUGH_DRAFT.docx.",
            "Editable draft: docs/dissertation/ETS_PHD_DISSERTATION_ROUGH_DRAFT.md.",
            "Advisor packet: docs/dissertation/MST_ADVISOR_REVIEW_PACKET.docx and docs/dissertation/MST_ADVISOR_REVIEW_PACKET.md.",
            "Sprint 2 audit artifacts: docs/dissertation/CLAIM_AUDIT.md, docs/dissertation/RESEARCH_ARTIFACT_MAP.md, and docs/dissertation/SPRINT_2_READINESS_REPORT.md.",
            "Sprint 3 literature artifacts: docs/dissertation/BIBLIOGRAPHY.md, docs/dissertation/RELATED_WORK_MATRIX.md, and docs/dissertation/SPRINT_3_READINESS_REPORT.md.",
            "Sprint 4 formal-methods artifacts: docs/dissertation/FORMAL_METHODS_AUDIT.md, docs/dissertation/PROOF_STATUS_TABLE.md, docs/dissertation/MODEL_CHECKING_COMMAND_LOG.md, and docs/dissertation/SPRINT_4_READINESS_REPORT.md.",
            "Sprint 5 implementation/reproducibility artifacts: docs/dissertation/IMPLEMENTATION_REPRODUCIBILITY_AUDIT.md, docs/dissertation/EXPERIMENT_ARTIFACT_PLAN.md, docs/dissertation/GOLDEN_VECTOR_COVERAGE.md, and docs/dissertation/SPRINT_5_READINESS_REPORT.md.",
            "Sprint 6 publication artifacts: docs/dissertation/PAPER_PIPELINE_ROADMAP.md, docs/dissertation/PAPER_ABSTRACTS_AND_OUTLINES.md, docs/dissertation/PAPER_CLAIM_EVIDENCE_MAP.md, docs/dissertation/VENUE_STRATEGY.md, and docs/dissertation/SPRINT_6_READINESS_REPORT.md.",
            "Sprint 7 dissertation assembly artifacts: docs/dissertation/DISSERTATION_ASSEMBLY_PLAN.md, docs/dissertation/CHAPTER_INTEGRATION_CHECKLIST.md, docs/dissertation/FIGURE_TABLE_PLAN.md, docs/dissertation/COMMITTEE_DRAFT_READINESS.md, and docs/dissertation/SPRINT_7_READINESS_REPORT.md.",
            "Sprint 8 defense-preparation artifacts: docs/dissertation/DEFENSE_QA.md, docs/dissertation/DEFENSE_DECK_PLAN.md, docs/dissertation/ARTIFACT_WALKTHROUGH_SCRIPT.md, docs/dissertation/FINAL_REVISION_CHECKLIST.md, and docs/dissertation/SPRINT_8_READINESS_REPORT.md.",
            "Relevant source folders: docs/dissertation, docs/research, formal/tla, formal/lean, tests/unit, ets/experiments, ets/benchmarks.",
            "Key official resources to confirm with staff: Missouri S&T Graduate Education, Missouri S&T Military & Veterans Service Center, Missouri S&T Student Financial Services, VA VR&E Chapter 31, VA GI Bill comparison tool, and Missouri Returning Heroes Act guidance.",
        ],
    ),
]


def build_markdown() -> str:
    lines = [
        f"# {TITLE}",
        "",
        f"## {SUBTITLE}",
        "",
        f"**Prepared for:** Advisor review / PhD return discussion",
        "",
        f"**Prepared by:** {AUTHOR}",
        "",
        f"**Date:** {DATE}",
        "",
        "**Status:** Draft packet for advisor review; not a formal university submission.",
        "",
    ]
    for heading, paragraphs in SECTIONS:
        lines.extend([f"# {heading}", ""])
        for paragraph in paragraphs:
            lines.extend([paragraph, ""])
    return "\n".join(lines).rstrip() + "\n"


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 14, 7),
        ("Heading 2", 13, "2E74B5", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "Packet Meta" not in doc.styles:
        meta = doc.styles.add_style("Packet Meta", WD_STYLE_TYPE.PARAGRAPH)
        meta.font.name = "Calibri"
        meta.font.size = Pt(10)
        meta.font.color.rgb = RGBColor(85, 85, 85)
        meta.paragraph_format.space_after = Pt(4)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(TITLE)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(SUBTITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    for label, value in [
        ("Prepared for", "Advisor review / PhD return discussion"),
        ("Prepared by", AUTHOR),
        ("Date", DATE),
        ("Status", "Draft packet for advisor review; not a formal university submission"),
    ]:
        p = doc.add_paragraph(style="Packet Meta")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    table.rows[0].cells[0].text = "Core ask"
    table.rows[0].cells[1].text = (
        "Review whether ETS can be shaped into a defensible Missouri S&T PhD "
        "dissertation and identify the administrative, funding, and research "
        "steps needed to return."
    )
    for cell in table.rows[0].cells:
        shade(cell, "F4F6F9")
    doc.add_page_break()


def build_docx() -> None:
    doc = Document()
    configure_doc(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = doc.styles["Packet Meta"]
    footer.add_run("MST Advisor Review Packet")
    add_cover(doc)

    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            if re.match(r"^\d+\.", paragraph):
                doc.add_paragraph(paragraph, style="List Number")
            else:
                doc.add_paragraph(paragraph)

    doc.save(DOCX_PATH)


def main() -> None:
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    words = len(re.findall(r"\b[\w'-]+\b", MD_PATH.read_text(encoding="utf-8")))
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Approximate packet word count: {words}")


if __name__ == "__main__":
    main()
